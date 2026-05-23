"""
MCP File Server
===============
Exposes basic file reading as an MCP tool.
"""
import os
from mcp.server.fastmcp import FastMCP
from loguru import logger

mcp = FastMCP("FileServer")

@mcp.tool()
def read_file(file_path: str) -> str:
    """
    Read the contents of a file. Supports TXT, MD, PDF, DOCX, and JSON formats.
    
    Args:
        file_path: Absolute path to the file.
        
    Returns:
        String content of the file.
    """
    if not os.path.exists(file_path):
        return f"Error: File not found at {file_path}"
        
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Reading file '{file_path}' with extension '{ext}'")
    
    try:
        if ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(f"--- Page {i+1} ---\n{page_text.strip()}")
                
                content = "\n\n".join(pages_text)
                if not content.strip():
                    content = "Warning: PDF parsed successfully, but no text could be extracted. It may be scanned/image-only."
                logger.info(f"Successfully read PDF: {file_path} ({len(reader.pages)} pages)")
                return content
            except ImportError:
                return "Error: The 'pypdf' package is required to read PDF files but is not installed."
                
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # Check for tables and extract their text too
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            table_texts.append(row_text)
                
                content = "\n\n".join(paragraphs)
                if table_texts:
                    content += "\n\n--- Tables in Document ---\n" + "\n".join(table_texts)
                    
                if not content.strip():
                    content = "Warning: Word document parsed successfully, but it contains no text paragraphs or tables."
                logger.info(f"Successfully read Word document: {file_path}")
                return content
            except ImportError:
                return "Error: The 'python-docx' package is required to read Word files but is not installed."
                
        elif ext == ".json":
            import json
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            content = json.dumps(data, indent=2, ensure_ascii=False)
            logger.info(f"Successfully read JSON file: {file_path}")
            return content
            
        else:
            # Default fallback: plain text UTF-8
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            logger.info(f"Successfully read text file: {file_path}")
            return content
            
    except Exception as e:
        logger.error(f"Failed to read file {file_path}: {e}")
        return f"Error reading file: {str(e)}"

if __name__ == "__main__":
    logger.info("Starting File MCP Server (stdio)...")
    mcp.run()
